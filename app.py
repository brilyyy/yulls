from flask import Flask, request, send_file, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import os
import tempfile

app = Flask(__name__)

# Ensure the upload folder exists
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def remove_nested_tables(table):
    for row in table.rows:
        for cell in row.cells:
            if cell.tables:
                for nested_table in cell.tables:
                    cell._element.remove(nested_table._element)

                if not cell.paragraphs or (len(cell.paragraphs) == 1 and not cell.paragraphs[0].text):
                    p = cell.add_paragraph()
                    r = p.add_run()
                    r._element.append(OxmlElement('w:br'))


def remove_extra_columns(table):
    if len(table.columns) > 2:
        tbl = table._tbl
        tr_elements = tbl.xpath('.//w:tr')

        for tr in tr_elements:
            tc_elements = tr.xpath('.//w:tc')
            for tc in tc_elements[2:]:
                tr.remove(tc)

        tblGrid = tbl.xpath('.//w:tblGrid')[0]
        gridCol_elements = tblGrid.xpath('.//w:gridCol')

        for gridCol in gridCol_elements:
            tblGrid.remove(gridCol)

        for _ in range(2):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(Inches(3).twips)))
            tblGrid.append(gridCol)


def add_new_column(table):
    tbl = table._tbl
    tr_elements = tbl.xpath('.//w:tr')

    for i, tr in enumerate(tr_elements):
        new_tc = OxmlElement('w:tc')
        p = OxmlElement('w:p')
        new_tc.append(p)

        if i == 0:
            run = OxmlElement('w:r')
            text = OxmlElement('w:t')
            text.text = "Tingkat penyelesaian status tl"
            run.append(text)
            p.append(run)

        tr.append(new_tc)

    tblGrid = tbl.xpath('.//w:tblGrid')[0]
    gridCol_elements = tblGrid.xpath('.//w:gridCol')

    while len(gridCol_elements) < 3:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(Inches(2).twips)))
        tblGrid.append(gridCol)
        gridCol_elements = tblGrid.xpath('.//w:gridCol')

    total_width = sum(int(col.get(qn('w:w'))) for col in gridCol_elements)
    new_width = int(total_width / 3)
    for col in gridCol_elements:
        col.set(qn('w:w'), str(new_width))


@app.route('/process_document', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            doc = Document(filepath)

            for table in doc.tables:
                remove_nested_tables(table)
                remove_extra_columns(table)
                add_new_column(table)

            output_filename = 'output_' + filename
            output_filepath = os.path.join(
                app.config['UPLOAD_FOLDER'], output_filename)
            doc.save(output_filepath)

            return send_file(output_filepath, as_attachment=True)
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            # Clean up temporary files
            if os.path.exists(filepath):
                os.remove(filepath)
            if os.path.exists(output_filepath):
                os.remove(output_filepath)
    else:
        return jsonify({"error": "Invalid file format. Please upload a .docx file"}), 400


if __name__ == '__main__':
    app.run(debug=True)
