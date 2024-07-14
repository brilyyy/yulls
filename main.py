import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QFileDialog, QMessageBox
from PyQt5.QtCore import Qt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def remove_nested_tables(table):
    for row in table.rows:
        for cell in row.cells:
            # Cek apakah ada tabel di dalam sel
            if cell.tables:
                # Hapus semua tabel dalam sel
                for nested_table in cell.tables:
                    cell._element.remove(nested_table._element)

                # Tambahkan paragraf kosong ke sel jika sel menjadi kosong
                if not cell.paragraphs or (len(cell.paragraphs) == 1 and not cell.paragraphs[0].text):
                    p = cell.add_paragraph()
                    r = p.add_run()
                    r._element.append(OxmlElement('w:br'))


def remove_extra_columns(table):
    # Periksa apakah tabel memiliki lebih dari 2 kolom
    if len(table.columns) > 2:
        # Dapatkan referensi ke elemen XML tabel
        tbl = table._tbl

        # Dapatkan semua baris tabel
        tr_elements = tbl.xpath('.//w:tr')

        for tr in tr_elements:
            # Dapatkan semua sel dalam baris
            tc_elements = tr.xpath('.//w:tc')

            # Hapus semua sel kecuali dua pertama
            for tc in tc_elements[2:]:
                tr.remove(tc)

        # Perbarui properti lebar tabel
        tblGrid = tbl.xpath('.//w:tblGrid')[0]
        gridCol_elements = tblGrid.xpath('.//w:gridCol')

        # Hapus semua elemen gridCol yang ada
        for gridCol in gridCol_elements:
            tblGrid.remove(gridCol)

        # Tambahkan dua elemen gridCol baru
        for _ in range(2):
            gridCol = OxmlElement('w:gridCol')
            # Setiap kolom lebar 3 inci
            gridCol.set(qn('w:w'), str(int(Inches(3).twips)))
            tblGrid.append(gridCol)


def add_new_column(table):
    # Dapatkan referensi ke elemen XML tabel
    tbl = table._tbl

    # Dapatkan semua baris tabel
    tr_elements = tbl.xpath('.//w:tr')

    for i, tr in enumerate(tr_elements):
        # Buat sel baru
        new_tc = OxmlElement('w:tc')

        # Tambahkan paragraf ke sel baru
        p = OxmlElement('w:p')
        new_tc.append(p)

        # Jika ini adalah baris pertama (header), tambahkan judul kolom
        if i == 0:
            run = OxmlElement('w:r')
            text = OxmlElement('w:t')
            text.text = "Tingkat penyelesaian status tl"
            run.append(text)
            p.append(run)

        # Tambahkan sel baru ke baris
        tr.append(new_tc)

    # Perbarui properti lebar tabel
    tblGrid = tbl.xpath('.//w:tblGrid')[0]
    gridCol_elements = tblGrid.xpath('.//w:gridCol')

    # Jika jumlah kolom kurang dari 3, tambahkan kolom baru
    while len(gridCol_elements) < 3:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(Inches(2).twips)))  # Lebar kolom 2 inci
        tblGrid.append(gridCol)
        gridCol_elements = tblGrid.xpath('.//w:gridCol')

    # Sesuaikan lebar kolom yang ada
    total_width = sum(int(col.get(qn('w:w'))) for col in gridCol_elements)
    new_width = int(total_width / 3)
    for col in gridCol_elements:
        col.set(qn('w:w'), str(new_width))


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Document Processor")

        button = QPushButton("Select & Process .docx File")
        button.clicked.connect(self.process_file)
        self.setCentralWidget(button)

    def process_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select a .docx file", "", "Word Documents (*.docx);;All Files (*)"
        )
        if not file_path:
            return

        try:
            doc = Document(file_path)
            # Proses setiap tabel dalam dokumen
            for table in doc.tables:
                remove_nested_tables(table)
                remove_extra_columns(table)
                add_new_column(table)

            save_path, _ = QFileDialog.getSaveFileName(
                self, "Save the processed file as", "", "Word Documents (*.docx);;All Files (*)"
            )
            if not save_path:
                return

            doc.save(save_path)
            QMessageBox.information(
                self, "Success", f"File processed and saved at:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred:\n{e}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
